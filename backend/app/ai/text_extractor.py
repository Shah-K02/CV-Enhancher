import re
import fitz  # PyMuPDF
from docx import Document

class TextExtractor:
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        text = []
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text.append(page.get_text())
        except Exception as e:
            raise ValueError(f"Error reading PDF: {e}")
        return TextExtractor.clean_text(" ".join(text))
        
    @staticmethod  
    def extract_from_docx(file_path: str) -> str:
        text = []
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {e}")
        return TextExtractor.clean_text(" ".join(text))
        
    @classmethod
    def extract(cls, file_path: str, filename: str) -> str:
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            return cls.extract_from_pdf(file_path)
        elif ext in ["doc", "docx"]:
            return cls.extract_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Please upload PDF or DOCX.")
            
    @staticmethod
    def clean_text(text: str) -> str:
        # Remove excessive whitespace, normalize line breaks
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
